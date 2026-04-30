from __future__ import annotations

import csv
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


TEXT_EXTENSIONS = {".txt", ".md", ".tex", ".py", ".m", ".json", ".csv"}


def sanitize_filename(filename: str) -> str:
    name = Path(filename).name.strip() or "upload"
    safe = []
    for char in name:
        if char.isalnum() or char in "._-()[] ":
            safe.append(char)
        else:
            safe.append("_")
    return "".join(safe)[:160]


def extract_text(path: Path, max_chars: int = 80_000) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path, max_chars)
        if suffix == ".docx":
            return _extract_docx(path, max_chars)
        if suffix in {".xlsx", ".xlsm"}:
            return _extract_xlsx(path, max_chars)
        if suffix == ".csv":
            return _extract_csv(path, max_chars)
        if suffix in TEXT_EXTENSIONS:
            return _read_text(path, max_chars)
    except Exception as exc:
        return f"[Text extraction failed for {path.name}: {exc}]"
    return f"[Unsupported file type for text extraction: {path.name}]"


def _read_text(path: Path, max_chars: int) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding, errors="replace")[:max_chars]
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")[:max_chars]


def _extract_pdf(path: Path, max_chars: int) -> str:
    if PdfReader is None:
        return "[PyPDF2 is not installed; PDF text extraction unavailable.]"
    reader = PdfReader(str(path))
    chunks: list[str] = []
    total = 0
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        chunk = f"\n\n--- page {index} ---\n{text}"
        chunks.append(chunk)
        total += len(chunk)
        if total >= max_chars:
            break
    return "".join(chunks)[:max_chars]


def _extract_docx(path: Path, max_chars: int) -> str:
    if docx is None:
        return "[python-docx is not installed; DOCX text extraction unavailable.]"
    document = docx.Document(str(path))
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)[:max_chars]


def _extract_csv(path: Path, max_chars: int) -> str:
    text = _read_text(path, max_chars)
    sample = text.splitlines()[:40]
    try:
        dialect = csv.Sniffer().sniff("\n".join(sample))
        rows = list(csv.reader(sample, dialect=dialect))
        preview = "\n".join(" | ".join(row[:12]) for row in rows[:20])
        return f"CSV preview:\n{preview}\n\nRaw excerpt:\n{text[:max_chars]}"
    except Exception:
        return text[:max_chars]


def _extract_xlsx(path: Path, max_chars: int) -> str:
    if openpyxl is None:
        return "[openpyxl is not installed; XLSX text extraction unavailable.]"
    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    chunks: list[str] = []
    total = 0
    for sheet in workbook.worksheets[:8]:
        chunks.append(f"\n\n--- sheet: {sheet.title} ---")
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            cells = ["" if value is None else str(value) for value in row[:20]]
            if any(cell.strip() for cell in cells):
                line = " | ".join(cells)
                chunks.append(line)
                total += len(line)
            if row_index >= 80 or total >= max_chars:
                break
        if total >= max_chars:
            break
    return "\n".join(chunks)[:max_chars]
